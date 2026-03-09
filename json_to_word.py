import json
import docx
import os

def formatear_tiempo_assembly(milisegundos: int) -> str:
    """Convierte los milisegundos de AssemblyAI a formato MM:SS."""
    segundos_totales = milisegundos // 1000
    minutos = int(segundos_totales // 60)
    segundos = int(segundos_totales % 60)
    return f"{minutos:02d}:{segundos:02d}"

def procesar_assembly_a_word(json_file: str, output_word_file: str):
    """
    Lee el JSON local de AssemblyAI y genera un documento Word estructurado.
    """
    print(f"Leyendo resultados de: {json_file}")
    
    try:
        with open(json_file, 'r', encoding='utf-8') as f:
            datos = json.load(f)
    except FileNotFoundError:
        print(f"[ERROR] No se encontró el archivo {json_file}")
        return
        
    doc = docx.Document()
    doc.add_heading('Transcripción Estructurada (AssemblyAI)', 0)

    # AssemblyAI guarda los diálogos separados en la lista 'utterances'
    utterances = datos.get('utterances', [])
    
    if not utterances:
        print("[ERROR] El JSON no contiene la lista de 'utterances'.")
        return

    print("Procesando estructura de hablantes y construyendo el documento...")

    for utterance in utterances:
        hablante = utterance.get('speaker', 'Desconocido')
        texto = utterance.get('text', '')
        tiempo_inicio_ms = utterance.get('start', 0)
        
        tiempo_formateado = formatear_tiempo_assembly(tiempo_inicio_ms)
        
        # Armamos el párrafo tal cual lo viste en tu terminal
        parrafo = f"Hablante {hablante} [{tiempo_formateado}]: {texto}"
        doc.add_paragraph(parrafo)

    doc.save(output_word_file)
    print("-" * 50)
    print(f"[ÉXITO] Documento Word generado correctamente: {output_word_file}")

if __name__ == "__main__":
    # Apuntamos al archivo JSON que te acaba de generar el script anterior
    ARCHIVO_JSON = "resultado_assembly.json" 
    
    # El nombre de tu documento final
    ARCHIVO_SALIDA_WORD = "Transcripcion_OLivia.docx"
    
    if os.path.exists(ARCHIVO_JSON):
        procesar_assembly_a_word(ARCHIVO_JSON, ARCHIVO_SALIDA_WORD)
    else:
        print(f"Por favor, asegúrate de que el archivo {ARCHIVO_JSON} esté en esta carpeta.")